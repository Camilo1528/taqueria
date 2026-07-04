import os
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed yet")
    exit(1)

def add_title(doc, text):
    p = doc.add_heading(text, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading1(doc, text):
    doc.add_heading(text, level=1)

def add_heading2(doc, text):
    doc.add_heading(text, level=2)

def add_bullet(doc, title, description):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ": ").bold = True
    p.add_run(description)

def generate_doc():
    document = Document()
    
    add_title(document, 'Manual y Caracteristicas Completas de Taqueria Pro')
    
    document.add_paragraph('Este documento contiene la especificacion tecnica y de negocio completa de la suite de software Taqueria Pro, abarcando todas sus funcionalidades operativas y administrativas actualizadas a su ultima version.')
    
    add_heading1(document, '1. Modulos de Operacion (Punto de Venta y Restaurante)')
    
    add_heading2(document, '1.1 Panel de Caja (Cashier POS)')
    add_bullet(document, 'Toma de Pedidos Rapida', 'Creacion de ordenes presenciales vinculadas a mesas o para llevar.')
    add_bullet(document, 'Gestion de Modificadores', 'Posibilidad de agregar notas e instrucciones especiales a los platillos (ej. "Sin cebolla", "Salsa aparte").')
    add_bullet(document, 'Cobro Completo y Division de Cuentas (Split)', 'Soporte para multiples metodos de pago, calculo automatico de vueltas, y capacidad de dividir la cuenta entre varias personas.')
    add_bullet(document, 'Gestion Inteligente de Propinas', 'Calculo automatico de propina sugerida (10%) en el panel de pago, agilizando el cobro y permitiendo al cajero modificar o eliminar el valor facilmente.')
    add_bullet(document, 'Gestion de Domicilios (Delivery)', 'Creacion de pedidos a domicilio con captura de datos del cliente (nombre, telefono, direccion) desde la caja.')
    add_bullet(document, 'Customer Facing Display (CFD)', 'Sincronizacion en tiempo real con una pantalla secundaria (WebSocket) para que el cliente vea los productos y el total mientras se le cobra.')
    
    add_heading2(document, '1.2 Panel de Meseros (Waiters)')
    add_bullet(document, 'Comandas Moviles', 'Interfaz optimizada para dispositivos moviles para tomar pedidos directamente en la mesa.')
    add_bullet(document, 'Mapeo de Mesas', 'Visualizacion del estado de las mesas (disponibles, ocupadas).')
    
    add_heading2(document, '1.3 Pantalla de Cocina (KDS - Kitchen Display System)')
    add_bullet(document, 'Sincronizacion en Tiempo Real', 'Los pedidos aparecen automaticamente en la cocina via WebSockets.')
    add_bullet(document, 'Control de Tiempos y Estados', 'Seguimiento de ordenes (Nuevas, En Preparacion, Listas) con alertas visuales de tiempo.')
    
    add_heading1(document, '2. Auto-Servicio y Experiencia del Cliente')
    add_heading2(document, '2.1 Pedidos desde la Mesa (QR)')
    add_bullet(document, 'Menu Interactivo', 'El cliente escanea el codigo QR de su mesa y accede a un menu web completo con imagenes atractivas.')
    add_bullet(document, 'Envio directo a cocina', 'El cliente arma su orden y la envia directamente sin llamar al mesero.')
    add_bullet(document, 'Pagos desde el movil', 'Los clientes pueden realizar pagos o enviar notificaciones de cobro directamente desde su interfaz web.')
    
    add_heading2(document, '2.2 Plataforma de Domicilios y Donaciones')
    add_bullet(document, 'Modo Domicilio Nativo', 'Los clientes pueden abrir el menu desde sus casas y realizar pedidos pidiendo automaticamente nombre, direccion y telefono en el carrito de compras.')
    add_bullet(document, 'Campanas de Donaciones', 'Permite crear mesas especiales etiquetadas como "Donacion" para recibir pagos sin afectar la operacion local.')
    
    add_heading1(document, '3. Gestion y Administracion (Back-Office)')
    
    add_heading2(document, '3.1 Dashboard y Analiticas')
    add_bullet(document, 'Metricas en Vivo', 'Ventas del dia, flujo de caja, desglose por metodos de pago y ordenes cerradas.')
    add_bullet(document, 'Reportes y Tendencias', 'Analisis de las ventas de los ultimos dias y productos estrella (Best-sellers).')
    
    add_heading2(document, '3.2 Inventario, Recetas e Inteligencia Artificial')
    add_bullet(document, 'Control de Insumos', 'Creacion de items de inventario (gramos, unidades, litros) con alertas de stock minimo.')
    add_bullet(document, 'Recetas (Escandallos)', 'Vinculacion de insumos a productos. Al vender un Taco, se descuentan automaticamente los gramos exactos de carne, tortillas y salsas.')
    add_bullet(document, 'Control de Mermas (Waste)', 'Registro de productos caducados o daados.')
    add_bullet(document, 'Predictivo de Compras (AI)', 'Algoritmo que analiza el consumo de la ultima semana cruzado con las recetas para recomendar automaticamente que cantidades de inventario comprar.')
    
    add_heading2(document, '3.3 Empleados y Control de Asistencia')
    add_bullet(document, 'Roles', 'Administrador, Cajero, Mesero y Cocinero.')
    add_bullet(document, 'Login por PIN y Codigo QR', 'Generacion de credenciales con Codigos QR para que los empleados registren sus ingresos (Log de Asistencia).')
    add_bullet(document, 'Asignacion de Domiciliarios', 'Los domicilios pueden ser asignados a empleados registrados especificos.')
    add_bullet(document, 'Bonos de Cumpleaos', 'El sistema detecta cuando es el cumpleaos de un empleado y permite emitir un bono de regalo (Gift Card) automatico.')
    
    add_heading2(document, '3.4 Configuracion e Imagen Corporativa')
    add_bullet(document, 'Control de Imagenes de Productos', 'Subida local de fotos de cada platillo directamente desde el dispositivo del administrador hacia el servidor.')
    add_bullet(document, 'Pantalla de Publicidad (CFD)', 'Gestion de un carrusel de imagenes promocionales y publicitarias que se muestran en la pantalla rotativa del cliente. Se cargan facilmente desde la seccion de Configuracion y cambian cada 10 segundos de forma dinamica.')
    
    add_heading1(document, '4. Fidelizacion, Marketing e Integraciones')
    
    add_heading2(document, '4.1 Fidelizacion (Gift Cards / Vouchers)')
    add_bullet(document, 'Tarjetas de Regalo', 'Emision de codigos promocionales o bonos de regalo con saldo prepago.')
    add_bullet(document, 'Redencion Parcial', 'El cliente puede pagar en caja con su Gift Card, y si el pedido es menor, conservar el saldo restante.')
    
    add_heading2(document, '4.2 Motor de Reservaciones')
    add_bullet(document, 'Modulo de Reservas', 'Captura de datos de reservas (nombre, fecha, hora, cantidad de personas) con control de estado (Pendiente, Confirmada, Cancelada).')
    
    add_heading2(document, '4.3 Integraciones Tributarias y Externas')
    add_bullet(document, 'Facturacion Electronica DIAN', 'Endpoint disenado para sincronizar las ordenes cerradas con la DIAN usando su API KEY, emitiendo comprobantes validos para Colombia.')
    add_bullet(document, 'Integracion Multi-Sucursal (Franquicias)', 'El sistema posee endpoints maestros (`/api/sync/master` y `/api/sync_receive`) para sincronizar datos financieros entre diferentes sedes (Branch IDs).')
    add_bullet(document, 'WhatsApp Integrado', 'Soporte para cargar y guardar el numero de WhatsApp del negocio y registrar logs estructurados para conexion con bots externos.')
    
    # NUEVA SECCION TECNICA
    add_heading1(document, '5. Anexo Tecnico: Lista de Servicios (API Endpoints)')
    doc_apis = [
        ("Auth", "POST /api/login, POST /api/login/pin"),
        ("Productos", "GET /menu, POST /api/products, PUT /api/products/{id}, DELETE /api/products/{id}"),
        ("Customer", "GET /api/customer/menu, POST /api/customer/order, POST /api/customer/pay/{id}"),
        ("Kitchen", "GET /kitchen_data, WS /ws/kitchen"),
        ("POS / Caja", "POST /api/pos/pay, GET /api/pos/orders, PATCH /api/pos/orders/{id}/items/{id}/notes, POST /add_to_order"),
        ("Asistencia", "POST /api/attendance/log"),
        ("Inventario", "GET /api/inventory, POST /api/inventory, POST /api/inventory/{id}/add_stock, POST /api/inventory/{id}/waste"),
        ("Recetas", "GET /api/products/{id}/recipe, POST /api/products/{id}/recipe, DELETE /api/products/{id}/recipe/{item_id}"),
        ("IA Predictiva", "GET /api/inventory/predict"),
        ("Usuarios", "GET /api/users, POST /api/users, DELETE /api/users/{id}"),
        ("Reportes", "GET /api/reports/sales_trend, GET /dashboard_data"),
        ("Reservaciones", "GET /api/reservations, POST /api/reservations, POST /api/reservations/{id}/status"),
        ("Domicilios", "GET /api/delivery/pending, POST /api/delivery/assign/{id}, POST /api/pos/delivery, POST /api/delivery/complete/{id}"),
        ("Gift Cards", "GET /api/vouchers, POST /api/vouchers/gift_card, GET /api/vouchers/{code}/validate, POST /api/vouchers/{code}/redeem"),
        ("Empleados Extras", "GET /api/employees/birthdays, POST /api/employees/issue_birthday_bonus"),
        ("Integraciones", "POST /api/dian/invoice, POST /api/sync/master, POST /api/sync_receive"),
        ("Archivos", "POST /api/upload_logo, POST /api/upload_image")
    ]
    for cat, apis in doc_apis:
        add_bullet(document, cat, apis)
    
    # Save
    filepath = os.path.join(os.path.dirname(__file__), "Caracteristicas_Taqueria.docx")
    document.save(filepath)
    print(f"Document saved to {filepath}")

if __name__ == "__main__":
    generate_doc()
